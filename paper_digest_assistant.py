import streamlit as st
import os
import json
import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader
from io import BytesIO
from openai import OpenAI
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def fetch_pdf_text(url: str, max_pages: int = 5) -> str:
    """
    Download a PDF from URL and extract text from the first max_pages pages.
    """
    try:
        response = requests.get(url, timeout=30, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        response.raise_for_status()
        
        pdf_file = BytesIO(response.content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        pages_to_read = min(max_pages, len(reader.pages))
        
        for i in range(pages_to_read):
            page = reader.pages[i]
            text_parts.append(page.extract_text())
        
        full_text = "\n\n".join(text_parts)
        return clean_text(full_text)
        
    except Exception as e:
        raise Exception(f"Failed to fetch PDF from URL: {str(e)}")


def fetch_html_text(url: str) -> str:
    """
    Fetch HTML content from URL and extract readable text using BeautifulSoup.
    """
    try:
        response = requests.get(url, timeout=30, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Extract text from paragraph tags
        paragraphs = soup.find_all('p')
        text_parts = [p.get_text() for p in paragraphs if p.get_text().strip()]
        
        # If no paragraphs found, try to get all text
        if not text_parts:
            text_parts = [soup.get_text()]
        
        full_text = "\n\n".join(text_parts)
        return clean_text(full_text)
        
    except Exception as e:
        raise Exception(f"Failed to fetch HTML from URL: {str(e)}")


def clean_text(text: str, max_length: int = 10000) -> str:
    """
    Clean and optionally truncate text to a safe length.
    """
    # Remove extra whitespace
    text = " ".join(text.split())
    text = text.strip()
    
    # Truncate if too long
    if len(text) > max_length:
        text = text[:max_length] + "... [truncated]"
    
    if not text or len(text.strip()) < 100:
        raise ValueError("Extracted text is too short or empty. The content may be inaccessible.")
    
    return text


def summarize_paper(text: str, api_key: str) -> dict:
    """
    Use OpenAI API to summarize the paper into structured JSON.
    Returns a dictionary with all required fields.
    """
    try:
        client = OpenAI(api_key=api_key)
        
        system_prompt = "You are an assistant that summarizes academic papers into structured JSON. Always respond with ONLY a valid JSON object, no additional text or markdown formatting."
        
        user_prompt = f"""Given the following academic text, extract the following fields and respond ONLY in valid JSON:

{{
    "title": "The title of the paper",
    "field_or_topic": "The academic field or topic area",
    "research_question": "The main research question(s) addressed",
    "methodology": "The research methodology used",
    "key_findings": ["Finding 1", "Finding 2", "Finding 3"],
    "limitations": "The limitations of the study",
    "personal_takeaway": "Why this paper matters for someone interested in marketing, data, and international/cross-border e-commerce"
}}

Text:
{text}

Return ONLY the JSON object, no markdown code blocks, no explanations, just the raw JSON."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        
        content = response.choices[0].message.content.strip()
        
        # Remove markdown code blocks if present
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        # Parse JSON
        try:
            summary = json.loads(content)
        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse JSON from OpenAI response: {str(e)}. Response was: {content[:200]}")
        
        # Validate required keys
        required_keys = ["title", "field_or_topic", "research_question", "methodology", 
                        "key_findings", "limitations", "personal_takeaway"]
        for key in required_keys:
            if key not in summary:
                raise Exception(f"Missing required key '{key}' in OpenAI response")
        
        # Ensure key_findings is a list
        if not isinstance(summary["key_findings"], list):
            if isinstance(summary["key_findings"], str):
                summary["key_findings"] = [summary["key_findings"]]
            else:
                summary["key_findings"] = []
        
        return summary
        
    except Exception as e:
        if "api_key" in str(e).lower() or "authentication" in str(e).lower() or "invalid" in str(e).lower():
            raise Exception(f"OpenAI API authentication failed: {str(e)}")
        raise Exception(f"Error summarizing paper with OpenAI: {str(e)}")


def save_to_word(summary: dict, source_type: str, source_value: str, filename: str = None) -> str:
    """
    Save paper summary to a Word document.
    Returns the filename of the saved document.
    """
    try:
        # Generate filename if not provided
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # Clean title for filename
            safe_title = "".join(c for c in summary["title"][:50] if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_title = safe_title.replace(' ', '_')
            filename = f"paper_digest_{safe_title}_{timestamp}.docx"
        
        # Create a new Document
        doc = Document()
        
        # Set document title
        title = doc.add_heading(summary["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add metadata section
        doc.add_heading('Document Information', level=1)
        
        metadata_para = doc.add_paragraph()
        metadata_para.add_run('Date: ').bold = True
        metadata_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        metadata_para.add_run('\n')
        metadata_para.add_run('Source Type: ').bold = True
        metadata_para.add_run(source_type)
        if source_type == "url":
            metadata_para.add_run('\n')
            metadata_para.add_run('Source URL: ').bold = True
            metadata_para.add_run(source_value)
        
        doc.add_paragraph()  # Empty line
        
        # Field or Topic
        doc.add_heading('Field or Topic', level=1)
        doc.add_paragraph(summary["field_or_topic"])
        doc.add_paragraph()
        
        # Research Question
        doc.add_heading('Research Question', level=1)
        doc.add_paragraph(summary["research_question"])
        doc.add_paragraph()
        
        # Methodology
        doc.add_heading('Methodology', level=1)
        doc.add_paragraph(summary["methodology"])
        doc.add_paragraph()
        
        # Key Findings
        doc.add_heading('Key Findings', level=1)
        for finding in summary["key_findings"]:
            para = doc.add_paragraph(finding, style='List Bullet')
        doc.add_paragraph()
        
        # Limitations
        doc.add_heading('Limitations', level=1)
        doc.add_paragraph(summary["limitations"])
        doc.add_paragraph()
        
        # Personal Takeaway
        doc.add_heading('Personal Takeaway', level=1)
        doc.add_paragraph(summary["personal_takeaway"])
        
        # Save document
        doc.save(filename)
        
        return filename
        
    except Exception as e:
        raise Exception(f"Failed to save Word document: {str(e)}")


def main():
    st.title("📚 Research Paper Digest Assistant")
    st.write("Automate your paper reading workflow: extract key information and create structured reading notes.")
    
    # OpenAI API Key input at the top
    st.subheader("🔑 OpenAI API Key")
    openai_api_key = st.text_input("Enter your OpenAI API Key", type="password", key="openai_key_input")
    if openai_api_key:
        st.session_state["openai_api_key"] = openai_api_key
    
    st.divider()
    
    # Input mode selection
    st.subheader("📄 Input Mode")
    mode = st.radio("Choose Input Mode:", ["URL", "Raw Text"], horizontal=True)
    
    paper_text = None
    source_value = None
    
    if mode == "URL":
        paper_url = st.text_input("Paper URL", key="paper_url_input")
        source_value = paper_url
        
        if paper_url:
            if st.button("Fetch Content", key="fetch_button"):
                try:
                    with st.spinner("Fetching content..."):
                        if paper_url.lower().endswith('.pdf') or 'pdf' in paper_url.lower():
                            paper_text = fetch_pdf_text(paper_url)
                            st.success(f"✅ Fetched PDF content ({len(paper_text)} characters)")
                        else:
                            paper_text = fetch_html_text(paper_url)
                            st.success(f"✅ Fetched HTML content ({len(paper_text)} characters)")
                        
                        st.session_state["paper_text"] = paper_text
                        st.session_state["source_type"] = "url"
                        st.session_state["source_value"] = paper_url
                except Exception as e:
                    st.error(f"❌ Error fetching content: {str(e)}")
            
            # Show fetched text if available
            if "paper_text" in st.session_state and st.session_state.get("source_value") == paper_url:
                paper_text = st.session_state["paper_text"]
                with st.expander("Preview fetched text"):
                    st.text(paper_text[:1000] + "..." if len(paper_text) > 1000 else paper_text)
    
    else:  # Raw Text mode
        paper_text = st.text_area("Paste paper text here", height=300, key="raw_text_input")
        source_value = "raw_text"
        if paper_text:
            st.session_state["paper_text"] = paper_text
            st.session_state["source_type"] = "raw_text"
            st.session_state["source_value"] = "raw_text"
    
    st.divider()
    
    # Generate Digest button
    if st.button("Generate Digest", type="primary", use_container_width=True):
        # Validate API key
        api_key = st.session_state.get("openai_api_key") or os.getenv("OPENAI_API_KEY")
        if not api_key:
            st.error("❌ No OpenAI API key provided. Please enter your key at the top.")
            st.stop()
        
        # Get paper text
        if mode == "URL":
            if "paper_text" not in st.session_state or not st.session_state.get("paper_text"):
                st.error("❌ Please fetch content from URL first by clicking 'Fetch Content'.")
                st.stop()
            paper_text = st.session_state["paper_text"]
        else:
            if not paper_text or len(paper_text.strip()) < 100:
                st.error("❌ Please paste paper text (at least 100 characters).")
                st.stop()
        
        try:
            with st.status("🔄 Generating digest...", expanded=True) as status:
                st.write("🤖 Summarizing with OpenAI...")
                summary = summarize_paper(paper_text, api_key)
                st.write("✅ Summary generated")
                
                st.write("💾 Saving to Word document...")
                saved_filename = save_to_word(
                    summary,
                    st.session_state.get("source_type", mode.lower()),
                    st.session_state.get("source_value", source_value or "unknown")
                )
                st.write(f"✅ Saved to {saved_filename}")
            
            st.success("✅ Digest generated and saved successfully!")
            
            st.divider()
            
            # Display formatted reading note
            st.subheader("📝 Reading Note")
            
            st.markdown(f"### {summary['title']}")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**Field or Topic:** {summary['field_or_topic']}")
            with col2:
                st.markdown(f"**Methodology:** {summary['methodology']}")
            
            st.markdown(f"**Research Question:** {summary['research_question']}")
            
            st.markdown("**Key Findings:**")
            for finding in summary['key_findings']:
                st.markdown(f"- {finding}")
            
            st.markdown(f"**Limitations:** {summary['limitations']}")
            
            st.markdown(f"**Personal Takeaway:** {summary['personal_takeaway']}")
            
            st.divider()
            
            # Collapsible raw JSON section
            with st.expander("📋 Raw JSON"):
                st.json(summary)
            
            st.info(f"💾 Saved to Word document: {saved_filename}")
            
            # Provide download button
            with open(saved_filename, "rb") as file:
                st.download_button(
                    label="📥 Download Word Document",
                    data=file,
                    file_name=saved_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
        except Exception as e:
            st.error(f"❌ An error occurred: {str(e)}")


if __name__ == "__main__":
    main()

